"""
Phase 1: Document extraction - Extract text from all source documents.
"""
import os
import json
from pypdf import PdfReader
import docx

DIR = r"C:\Users\trangtt1\HH.new"
intel_dir = os.path.join(DIR, "intel")

docs = [
    ("1.TKCT-Cuc HH-KetCauHaTangGiaoThong 13.11.2024 ban in final.docx", "docx", "TKCT"),
    ("MOT_VMD_MTIS_DD_ 3.0_PHCV_Final (2).pdf", "pdf", "DD"),
    ("MOT_VMD_MTIS_TH_1.0 (2).pdf", "pdf", "TH"),
    ("MOT_VMD_MTIS_URD_ v3.0_PHCV_Final (1).pdf", "pdf", "URD"),
    ("MOT_VMD_MTIS_URD_v3.0_PHCV (3).docx", "docx", "URD-docx"),
    ("VMD_MTIS_TaiLieuKhaoSat_PM_NoiBo_v1.0_PHCV.pdf", "pdf", "Survey"),
    ("VMD_MTIS_TaiLieuKhaoSat_PM_NoiBo_v1.0_Phụ lục_PHCV.pdf", "pdf", "Survey-Appendix"),
]

extraction_report = []

for name, fmt, kind in docs:
    path = os.path.join(DIR, name)
    size = os.path.getsize(path)
    text = ""
    
    try:
        if fmt == "pdf":
            reader = PdfReader(path)
            pages = len(reader.pages)
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text += f"\n--- Page {i+1} ---\n{page_text}"
            extraction_report.append({
                "filename": name,
                "format": "PDF",
                "kind": kind,
                "size_bytes": size,
                "size_mb": round(size / (1024*1024), 1),
                "pages": pages,
                "text_length": len(text),
                "text_preview": text[:2000].replace('\n', ' '),
                "status": "extracted"
            })
        elif fmt == "docx":
            doc = docx.Document(path)
            pages_approx = max(1, len(doc.paragraphs) // 50)
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"
            # Also extract tables
            for table_idx, table in enumerate(doc.tables):
                text += f"\n[Table {table_idx+1}]\n"
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text += row_text + "\n"
            extraction_report.append({
                "filename": name,
                "format": "DOCX",
                "kind": kind,
                "size_bytes": size,
                "size_mb": round(size / (1024*1024), 1),
                "pages_approx": pages_approx,
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "text_length": len(text),
                "text_preview": text[:2000].replace('\n', ' '),
                "status": "extracted"
            })
    except Exception as e:
        extraction_report.append({
            "filename": name,
            "format": fmt.upper(),
            "kind": kind,
            "size_bytes": size,
            "status": f"error: {str(e)}"
        })

# Save extraction report
with open(os.path.join(intel_dir, "01-extraction-report.json"), "w", encoding="utf-8") as f:
    json.dump(extraction_report, f, ensure_ascii=False, indent=2)

print(f"Extraction complete: {len(extraction_report)} documents processed")
for r in extraction_report:
    print(f"  {r['filename']}: {r['size_mb']} MB, {r['pages'] if 'pages' in r else r.get('paragraphs','?')} pages/paras, text={r['text_length']} chars")
